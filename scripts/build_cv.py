from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "public/Omar_Tarboush_CV.docx"
PORTFOLIO_URL = "https://omartarboush.github.io/portfolio/"
PORTFOLIO_LABEL = "omartarboush/portfolio"
LINKEDIN_URL = "https://sy.linkedin.com/in/omar-tarboush"
LINKEDIN_LABEL = "Omar Tarboush"

INK = RGBColor(17, 24, 39)
MUTED = RGBColor(82, 92, 108)
ACCENT = RGBColor(11, 117, 209)


def set_run(run, size=8.7, bold=False, color=INK):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_text(paragraph, text, size=8.7, bold=False, color=INK):
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return run


def add_hyperlink(paragraph, text, url, size=8.7, bold=False):
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")

    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Arial")
    font.set(qn("w:hAnsi"), "Arial")
    run_props.append(font)

    size_element = OxmlElement("w:sz")
    size_element.set(qn("w:val"), str(int(size * 2)))
    run_props.append(size_element)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0B75D1")
    run_props.append(color)

    if bold:
        bold_element = OxmlElement("w:b")
        run_props.append(bold_element)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(underline)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(run_props)
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    add_text(p, text.upper(), size=9.25, bold=True, color=ACCENT)
    return p


def add_compact_paragraph(doc, text, size=8.15):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.2)
    p.paragraph_format.line_spacing = 1.05
    add_text(p, text, size=size, color=INK)
    return p


def add_labeled_line(doc, label, value, size=8.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.1)
    p.paragraph_format.line_spacing = 1.05
    add_text(p, f"{label}: ", size=size, bold=True, color=ACCENT)
    add_text(p, value, size=size, color=INK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(1.0)
    p.paragraph_format.line_spacing = 1.08
    add_text(p, text, size=8.65, color=INK)


def add_role(doc, title, company, location, period, bullets):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5.6)
    p.paragraph_format.space_after = Pt(0.9)
    p.paragraph_format.keep_with_next = True
    add_text(p, title, size=9.35, bold=True, color=INK)
    add_text(p, f" | {company}", size=9.35, bold=True, color=ACCENT)
    add_text(p, f" | {location}", size=8.45, color=MUTED)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(1.7)
    p2.paragraph_format.keep_with_next = True
    add_text(p2, period, size=8.25, bold=True, color=MUTED)

    for bullet in bullets:
        add_bullet(doc, bullet)


def add_project(doc, name, role, summary, stack, links, links_label="Stores"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2.8)
    p.paragraph_format.space_after = Pt(0.75)
    p.paragraph_format.keep_with_next = True
    add_text(p, name, size=8.85, bold=True, color=INK)
    add_text(p, f" | {role}", size=8.6, bold=True, color=ACCENT)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(0.75)
    p2.paragraph_format.line_spacing = 1.05
    add_text(p2, summary, size=8.2, color=INK)

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(1.4)
    add_text(p3, f"Stack: {stack}", size=7.9, bold=True, color=MUTED)
    add_text(p3, f" | {links_label}: {links}", size=7.9, color=MUTED)


def build():
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(0.34)
    section.bottom_margin = Inches(0.34)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(1.2)
    normal.paragraph_format.line_spacing = 1.05

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    add_text(title, "Omar Tarboush", size=19, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(1)
    add_text(
        subtitle,
        "Flutter Mobile & Web Developer | AI Graduate",
        size=9.8,
        bold=True,
        color=ACCENT,
    )

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(2.5)
    add_text(
        contact,
        "Damascus, Syria | +963 941 717 409 | omartarboush191@gmail.com | LinkedIn: ",
        size=7.6,
        bold=True,
        color=MUTED,
    )
    add_hyperlink(contact, LINKEDIN_LABEL, LINKEDIN_URL, size=7.6, bold=True)
    add_text(contact, " | Portfolio: ", size=7.6, bold=True, color=MUTED)
    add_hyperlink(contact, PORTFOLIO_LABEL, PORTFOLIO_URL, size=7.6, bold=True)

    add_section_heading(doc, "Professional Summary")
    add_compact_paragraph(
        doc,
        "Flutter developer and AI graduate with 4+ years of production mobile app experience and 12+ applications across the UAE, Saudi Arabia, and Syria. Strong in Flutter, Firebase, GetX, Riverpod, BLoC, REST/GraphQL APIs, Google Maps, deep links, notifications, secure access, and Patrol automated app testing. Delivered and maintained published apps across communication, food, beauty, task management, accounting, and location-based products, plus bilingual web delivery with Next.js.",
        size=8.05,
    )

    add_section_heading(doc, "Education & Languages")
    add_labeled_line(doc, "Education", "BSc Information Technology, AI Major - Arab International University", size=7.9)
    add_labeled_line(doc, "Languages", "Arabic native | English good", size=7.9)

    add_section_heading(doc, "Core Skills")
    skill_rows = [
        ("Mobile", "Flutter, Dart, GetX, Riverpod, BLoC, Patrol tests, Mockito, widget and integration testing"),
        ("Web", "Next.js, TypeScript, React, Prisma, Tailwind CSS, next-intl (i18n), SEO"),
        ("Backend/Data", "Firebase Auth, Firestore, Messaging, Cloud Functions, SQL, stored procedures, views"),
        ("Integrations", "REST APIs, GraphQL APIs, Google Maps, deep linking, push notifications, in-app purchases"),
        ("Security/Delivery", "Patrol automated app testing, PIN access, role-based access, data encryption, cloud backup, caching, Git, GitLab, GitHub"),
    ]
    for label, value in skill_rows:
        add_labeled_line(doc, label, value, size=7.75)

    add_section_heading(doc, "Work Experience")
    add_role(
        doc,
        "Full-Time Flutter Developer",
        "We Got Nerds",
        "Dubai, United Arab Emirates",
        "Oct 2025 - Present",
        [
            "Managed project timelines, delegated development tasks, and coordinated cross-functional delivery.",
            "Implemented Riverpod state management and optimized data queries for stronger app performance.",
            "Integrated Firebase Auth, Firestore, Analytics, third-party REST APIs, Patrol automated app tests, widget tests, and integration tests.",
        ],
    )
    add_role(
        doc,
        "Freelance Web Developer",
        "Freelance",
        "Remote",
        "2025 - Present",
        [
            "Designed and delivered a bilingual Arabic/English photographer website (Ghassan Ahmad) with seasonal offers, work galleries, WhatsApp booking, and localized SEO.",
            "Built the full web stack with Next.js, TypeScript, and Prisma, including a database-backed admin dashboard with image uploads and per-locale SEO control.",
        ],
    )
    add_role(
        doc,
        "Full-Time Flutter Developer",
        "ANS Soft",
        "Abu Dhabi, United Arab Emirates",
        "Aug 2023 - Oct 2025",
        [
            "Delivered and supported production mobile applications including Aklat24 and Mapy.",
            "Led special projects, planned schedules, supervised delivery work, and maintained large-scale legacy codebases.",
            "Built GetX architectures, SQL views, Firebase Messaging, deep links, GraphQL, REST APIs, and Patrol end-to-end automation tests.",
        ],
    )
    add_role(
        doc,
        "Full-Time Flutter Developer / Programming Instructor",
        "Seba School",
        "Damascus, Syria",
        "Sep 2023 - Sep 2024",
        [
            "Taught programming fundamentals with a focus on Python and practical problem solving.",
            "Organized student competitions and learning events to encourage hands-on practice.",
        ],
    )
    add_role(
        doc,
        "Full-Time Flutter Developer",
        "Goma Plus",
        "Saudi Arabia",
        "May 2023 - Aug 2023",
        [
            "Developed mobile apps using GetX and Firebase notifications.",
            "Created stored procedures, database views, complex queries, and REST API integrations.",
        ],
    )

    add_section_heading(doc, "Selected Projects")
    add_project(
        doc,
        "H1 SMS",
        "Unified SMS Manager",
        "Protected SMS workspace for multiple numbers with cloud backup, PIN access, encryption, and organized message workflows.",
        "Flutter, Firebase, Encryption, Cloud backup",
        "Play Store, App Store",
    )
    add_project(
        doc,
        "H1 Mail",
        "Secure Multi-Account Email",
        "Protected multi-mailbox email app with PIN login, encryption, mailbox restrictions, notification flows, and account organization.",
        "Flutter, IMAP, Firebase, Security",
        "Play Store, App Store",
    )
    add_project(
        doc,
        "H1 Assignment",
        "Team Task Platform",
        "Operational team platform covering tasks, groups, meetings, role-based access, profiles, real-time updates, and search.",
        "Flutter, Firebase, RBAC, Real-time",
        "Play Store, App Store",
    )
    add_project(
        doc,
        "Aklat24",
        "Meal-Sharing Marketplace",
        "Meal-sharing marketplace connecting meal seekers with providers to reduce food waste and support location-based discovery.",
        "Flutter, GetX, Maps, REST APIs",
        "Play Store, App Store",
    )
    add_project(
        doc,
        "Mapy",
        "Social Location App",
        "Location-based social app with Instagram-like sharing, hidden gems discovery, maps, profile flows, and content exploration.",
        "Flutter, Google Maps, GraphQL, Firebase",
        "Play Store, App Store",
    )
    add_project(
        doc,
        "Wonder Beauties",
        "Beauty Consultations",
        "Beauty consultation and commerce app with pharmacist/doctor consultations, personalized recommendations, and original product discovery.",
        "Flutter, Firebase, E-commerce, Consultations",
        "Play Store, App Store",
    )
    # Jeebtak hidden for now — restore by uncommenting this block.
    # add_project(
    #     doc,
    #     "Jeebtak",
    #     "Shop Accounting App",
    #     "Offline-first shop accounting product with multi-currency wallets, receipts, debt tracking, QR flows, PDF/Excel reports, SQLCipher-encrypted storage, and PIN/biometric access.",
    #     "Flutter, GetX, Drift, SQLCipher, Firebase",
    #     "Pre-launch (Android, iOS)",
    #     links_label="Status",
    # )
    add_project(
        doc,
        "Ghassan Ahmad",
        "Photographer Website",
        "Bilingual Arabic/English photographer website with seasonal offers, work galleries, WhatsApp booking, localized SEO, and a database-backed admin dashboard with image uploads.",
        "Next.js, TypeScript, Prisma, i18n",
        "Web (landing site + admin dashboard)",
        links_label="Platform",
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
